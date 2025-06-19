import requests

from fastapi import APIRouter, HTTPException, Depends, Request
from fastapi import UploadFile, File, Form
from fastapi.responses import JSONResponse
from fastapi.responses import PlainTextResponse

from sqlalchemy.orm import Session
from database import AsyncSessionLocal
from app.Utils.sendgrid import alert_courts_admin
from app.Model.CaseModel import TimeRange
from app.Model.CaseModel import FilterCondition
from app.Model.CaseModel import AlertAdminData
from app.Model.CaseModel import  UserNameModel
from app.Model.CaseModel import  ShortcodeModel
from app.Model.MainTable import TemplateModel
from app.Model.MainTable import TemplateCaseModel

import app.Utils.database_handler as crud
import app.Utils.Auth as Auth
from typing import Annotated
from dotenv import load_dotenv

from bs4 import BeautifulSoup
from docx import Document
from typing import List
import shutil
import os
import logging
import traceback

logger = logging.getLogger(__name__)

load_dotenv()
router = APIRouter()

UPLOAD_DIR = "upload"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Dependency to get the database session
async def get_db():
    async with AsyncSessionLocal() as session:
        yield session

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

@router.post("/upload")
# async def upload( letterFiles: List[UploadFile] = File(...), envelopeFiles: List[UploadFile] = File(...), username: str = Form(...), db: Session = Depends(get_db)):
async def upload( letterFiles: List[UploadFile] = File(...), username: str = Form(...), db: Session = Depends(get_db)):
    results = []

    try:
        # Create user-specific subfolder
        user_folder = os.path.join(UPLOAD_DIR, username)
        os.makedirs(user_folder, exist_ok=True)

        # --- Letter Templates ---
        for letterFile in letterFiles:
            ext = os.path.splitext(letterFile.filename)[1].lower()
            origin_name = letterFile.filename
            saved_name = origin_name  # Now identical
            saved_path = os.path.join(user_folder, saved_name)

            # Save file
            with open(saved_path, "wb") as buffer:
                shutil.copyfileobj(letterFile.file, buffer)

            # Extract human-readable content for response
            if ext == ".txt":
                with open(saved_path, "r", encoding="utf-8") as f:
                    readable_content = f.read()
            elif ext == ".docx":
                readable_content = read_docx(saved_path)
            else:
                readable_content = "Unsupported file type"

            # Read binary content for DB
            with open(saved_path, "rb") as f:
                binary_content = f.read()

            results.append({
                "filename": origin_name,
                "saved_as": saved_name,
                "content": readable_content
            })

            # Insert into DB
            letter_data = TemplateModel(
                origin_name=origin_name,
                saved_name=saved_name,
                saved_path=saved_path,
                template_type="letter",
                content=binary_content,
                user=username
            )
            await crud.insert_template(db, letter_data)

        # --- Envelope Templates ---
        # for envelopeFile in envelopeFiles:
        #     ext = os.path.splitext(envelopeFile.filename)[1].lower()
        #     origin_name = envelopeFile.filename
        #     saved_name = origin_name  # Now identical
        #     saved_path = os.path.join(user_folder, saved_name)

        #     # Save file
        #     with open(saved_path, "wb") as buffer:
        #         shutil.copyfileobj(envelopeFile.file, buffer)

        #     # Extract human-readable content for response
        #     if ext == ".txt":
        #         with open(saved_path, "r", encoding="utf-8") as f:
        #             readable_content = f.read()
        #     elif ext == ".docx":
        #         readable_content = read_docx(saved_path)
        #     else:
        #         readable_content = "Unsupported file type"

        #     # Read binary content for DB
        #     with open(saved_path, "rb") as f:
        #         binary_content = f.read()

        #     results.append({
        #         "filename": origin_name,
        #         "saved_as": saved_name,
        #         "content": readable_content
        #     })

        #     # Insert into DB
        #     envelope_data = TemplateModel(
        #         origin_name=origin_name,
        #         saved_name=saved_name,
        #         saved_path=saved_path,
        #         template_type="envelope",
        #         content=binary_content,
        #         user=username
        #     )

        #     await crud.insert_template(db, envelope_data)
            
        return {"success": True, "username": username, "files": results}

    except Exception as e:
        print("Error:", str(e))
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))
    
@router.post("/getCases")
async def getCases(timeRange: TimeRange, db: Session = Depends(get_db)):
    try:
        cases = await crud.get_cases(db, timeRange)
        return {"cases": cases}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/getCounties")
async def getCounties(timeRange: TimeRange, db: Session = Depends(get_db)):
    try:
        counties = await crud.get_counties(db, timeRange)
        return {"counties": counties}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/getData")
async def getData(filterCondition: FilterCondition, db: Session = Depends(get_db)):
    try:
        cases = await crud.get_data(db, filterCondition)
        return {"cases": cases}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/getDataForMerge")
async def getDataForMerge(filterCondition: FilterCondition, db: Session = Depends(get_db)):
    try:
        print("getDataForMerge:")
        cases = await crud.get_data_merge(db, filterCondition)
        return {"cases": cases}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/getLastQueryDate")
async def getLastQueryDate(db: Session = Depends(get_db)):
    try:
        queryDate = await crud.get_last_query_date(db)
        return {"query_date": queryDate}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/fetchCourts")
async def fetchCourts(db: Session = Depends(get_db)):
    try:
        url = "https://www.in.gov/courts/help/odyssey-courts/"

        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            tables = soup.find_all('table')

            if len(tables) >= 2:
                second_table = tables[1]  # Get the second table
                data = []
                for row in second_table.find_all('tr'):
                    cols = row.find_all('td')
                    if cols:
                        data.append([col.get_text(strip=True) for col in cols])       
                resultFlag = await crud.insert_courts(db, data)
                return {"success": resultFlag}
            else:
                print("The second table was not found on the page.")
                return {"success": False}
        else:
            print(f"Failed to retrieve data. Status code: {response.status_code}")
            return {"success": False}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/fetchCounties")
async def fetchCounties(db: Session = Depends(get_db)):
    try:
        url = "https://www.in.gov/courts/local/"

        response = requests.get(url)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')

            section = soup.find("section", id="645676")

            counties = []
            for a in section.find_all("a", href=True):
                county_name = a.text.strip()
                counties.append({"name": county_name})
                
            resultFlag = await crud.insert_counties(db, counties)
            return {"success": resultFlag}
        
        else:
            print(f"Failed to retrieve data. Status code: {response.status_code}")
            return {"success": False}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@router.post("/getCourts")
async def getCourts(db: Session = Depends(get_db)):
    try:
        courts = await crud.get_courts(db)
        return {"courts": courts}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@router.post("/getPaidCourts")
async def getPaidCourts(data: UserNameModel, db: Session = Depends(get_db)):
    try:
        username = data.username
        paid_courts = await crud.get_paid_courts(db, username)
        print("templates - paid_courts: ", paid_courts)
        return {"paid_courts": paid_courts}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/getPaidCounty")
async def getPaidCounty(data: UserNameModel, db: Session = Depends(get_db)):
    try:
        username = data.username
        paid_counties = await crud.get_paid_county(db, username)
        return {"paid_counties": paid_counties}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/getIndianaCounties")
async def getIndianaCounties(db: Session = Depends(get_db)):
    try:
        counties = await crud.get_indiana_counties(db)
        return {"counties": counties}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/alertCourtsToAdmin")
async def alertCourtsToAdmin(alertAdminData: AlertAdminData, db: Session = Depends(get_db)):
    try:
        await alert_courts_admin(alertAdminData, db)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@router.post("/getSavedTemplates")
async def getSavedTemplates(data: UserNameModel, db: Session = Depends(get_db)):
    try:
        username = data.username
        templates = await crud.get_templates(db, username)
        return {"templates": templates}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/getSavedShortcode")
async def getSavedShortcode(db: Session = Depends(get_db)):
    try:
        shortcodes = await crud.get_saved_shortcode(db)
        return {"shortcodes": shortcodes}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/getFields")
async def getFields(db: Session = Depends(get_db)):
    try:
        fields = await crud.get_fields(db)
        return {"fields": fields}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@router.post("/addNewShortcode")
async def addNewShortcode(shortcodes: ShortcodeModel, db: Session = Depends(get_db)):
    try:
        await crud.add_new_shortcode(db, shortcodes)
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/removeShortcode")
async def removeShortcode(shortcodes: ShortcodeModel, db: Session = Depends(get_db)):
    try:
        await crud.remove_shortcode(db, shortcodes)
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/removeSavedTemplate")
async def removeSavedTemplate(templates: TemplateModel, db: Session = Depends(get_db)):
    try:
        await crud.remove_saved_templates(db, templates)
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/getTemplateContent")
async def getTemplateContent(template: TemplateModel, db: Session = Depends(get_db)):
    try:
        saved_path = template.saved_path
        # templates = await crud.get_template_content(db, saved_path)
        print("templates - saved_path: ", saved_path)
        saved_path = saved_path.replace("\\", "/");
        if not os.path.exists(saved_path):
            return {"success" : False, "content": ''}
        ext = os.path.splitext(saved_path)[1].lower()
        if ext == ".txt":
            with open(saved_path, "r", encoding="utf-8") as f:
                return {"success" : True, "content": PlainTextResponse(f.read())}
        elif ext == ".docx":
            doc = Document(saved_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return {"success" : True, "content": PlainTextResponse(text)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/getCompletedTemplate")
async def getCompletedTemplate(data: TemplateCaseModel, db: Session = Depends(get_db)):
    filled_template = await crud.get_completed_template(db, data)
    return {"filled_template": filled_template}

@router.post("/getPurchasedCourts")
async def getPurchasedCourts(db: Session = Depends(get_db)):
    try:
        purchased_courts = await crud.get_purchased_courts(db)
        print("templates - purchased_courts: ", purchased_courts)
        return {"purchased_courts": purchased_courts}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@router.post("/getCountiesAllData")
async def getCountiesAllData(db: Session = Depends(get_db)):
    try:
        all_data = await crud.get_counties_all_data(db)
        return {"all_data": all_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
